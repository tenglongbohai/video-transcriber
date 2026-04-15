#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
视频转文字工具 (Video Transcriber)
支持: 视频 -> 音频提取 -> Whisper CUDA 转录 -> Word 文档
"""

import os
import sys
import json
import time
import queue
import signal
import subprocess
import traceback
from datetime import datetime
from pathlib import Path

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLineEdit, QPushButton, QLabel, QProgressBar, QTextEdit,
    QFileDialog, QComboBox, QCheckBox, QMessageBox, QGroupBox,
    QGridLayout, QStatusBar, QScrollArea, QFrame, QDialog
)
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QTimer
from PyQt5.QtGui import QFont, QTextCursor, QColor

import psutil
import requests


# ============================================================================
# 核心模块
# ============================================================================

class AudioExtractor:
    """FFmpeg 音频提取器"""

    def __init__(self, ffmpeg_path="ffmpeg"):
        self.ffmpeg_path = ffmpeg_path

    def extract_audio(self, video_path, output_path, sample_rate=16000, callback=None):
        """提取音频为 wav 格式"""
        cmd = [
            self.ffmpeg_path, "-y", "-i", video_path,
            "-vn", "-acodec", "pcm_s16le",
            "-ar", str(sample_rate), "-ac", "1",
            "-loglevel", "error", output_path
        ]

        try:
            process = subprocess.Popen(
                cmd, stdout=subprocess.PIPE,
                stderr=subprocess.PIPE, universal_newlines=True
            )

            # 进度监控
            duration = self._get_duration(video_path)
            while True:
                line = process.stderr.readline()
                if not line and process.poll() is not None:
                    break
                if line and "time=" in line:
                    current_time = self._parse_time(line)
                    if duration > 0 and callback:
                        progress = min(95, int(current_time / duration * 100))
                        callback(progress)

            if process.returncode == 0:
                if callback:
                    callback(100)
                return True, "音频提取完成"
            else:
                return False, f"FFmpeg 错误: {process.stderr.read()}"

        except Exception as e:
            return False, f"音频提取失败: {str(e)}"

    def _get_duration(self, video_path):
        """获取视频时长（秒）"""
        cmd = [self.ffmpeg_path, "-i", video_path]
        try:
            result = subprocess.run(
                cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE,
                universal_newlines=True, timeout=30
            )
            output = result.stderr
            for line in output.split('\n'):
                if "Duration:" in line:
                    try:
                        duration_str = line.split("Duration:")[1].split(",")[0].strip()
                        h, m, s = duration_str.split(":")
                        return int(h) * 3600 + int(m) * 60 + float(s)
                    except:
                        pass
        except:
            pass
        return 0

    def _parse_time(self, line):
        """解析 ffmpeg 时间输出"""
        try:
            time_str = line.split("time=")[1].split(" ")[0]
            h, m, s = time_str.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        except:
            return 0

    def get_audio_duration(self, audio_path):
        """获取音频时长"""
        return self._get_duration(audio_path)


class Transcriber:
    """Faster-Whisper 转录器"""

    def __init__(self, model_size="small", device="cuda", compute_type="default"):
        self.model_size = model_size
        self.device = device
        self.compute_type = compute_type
        self.model = None
        self.load_model()

    def load_model(self):
        """加载 Whisper 模型"""
        try:
            from faster_whisper import WhisperModel

            # 量化类型映射
            compute_type_map = {
                "float16": "float16",      # 高质量，显存占用高
                "int8": "int8",            # 量化，内存减半
                "int8_float16": "int8_float16",  # 混合量化，推荐
                "default": "float16" if self.device == "cuda" else "int8"
            }

            actual_compute = compute_type_map.get(self.compute_type, "float16")

            self.model = WhisperModel(
                self.model_size, device=self.device,
                compute_type=actual_compute
            )
            return True, f"模型加载成功 ({self.model_size}, {actual_compute})"
        except ImportError:
            return False, "请安装 faster-whisper: pip install faster-whisper"
        except Exception as e:
            return False, f"模型加载失败: {str(e)}"

    def transcribe(self, audio_path, chunk_start=0, chunk_duration=None, callback=None):
        """转录音频"""
        if self.model is None:
            success, msg = self.load_model()
            if not success:
                return False, msg, ""

        try:
            kwargs = {
                "audio": audio_path,
                "vad_filter": True,
                "language": "zh",
            }
            if chunk_duration:
                kwargs["max_new_tokens"] = 512

            segments, info = self.model.transcribe(**kwargs)

            text_parts = []
            segment_count = 0

            for segment in segments:
                # 计算当前片段的相对起始时间
                segment_start = segment.start - chunk_start if chunk_start else segment.start

                # 只处理指定时长范围内的片段
                if chunk_duration and segment_start >= chunk_duration:
                    break

                text = segment.text.strip()
                if text:
                    text_parts.append(text)

                segment_count += 1
                if callback and segment_count % 10 == 0:
                    # 估算进度（基于音频时长）
                    total_duration = info.duration or 1
                    progress = min(100, int(segment.end / total_duration * 100))
                    callback(progress, text)

            full_text = " ".join(text_parts)
            return True, "转录完成", full_text

        except Exception as e:
            return False, f"转录失败: {str(e)}", ""


class PunctuationAdder:
    """标点符号添加器"""

    def __init__(self):
        self.punctuation_rules = [
            (r'([。！？\?!])', r'\1'),  # 保留已有句末标点
            (r'([^。！？\?!\,，、])([，。、])', r'\1\2'),  # 添加逗号
            (r'(\d{4})年', r'\1年'),  # 年份格式
            (r'(\d+)月(\d+)日', r'\1月\2日'),  # 日期格式
        ]

    def add_punctuation(self, text):
        """为文本添加标点符号"""
        if not text:
            return text

        # 使用基础规则进行标点处理
        result = text

        # 在适当位置添加句号（基于句子长度和上下文）
        sentences = []
        current = ""

        for char in result:
            current += char
            if char in '。！？':
                sentences.append(current.strip())
                current = ""

        if current.strip():
            # 如果剩余文本较长，添加句号
            if len(current) > 10:
                sentences.append(current.strip() + "。")
            else:
                sentences.append(current.strip())

        # 合并句子
        final_text = " ".join(sentences) if sentences else result

        return final_text

    def split_paragraphs(self, text, max_length=500):
        """智能分段"""
        if not text:
            return []

        paragraphs = []
        current = ""
        sentences = []

        # 按句子分割
        import re
        sentence_endings = r'[。！？\?!]+'
        parts = re.split(sentence_endings, text)

        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue

            # 添加标点
            if i < len(parts) - 1 or text[-1] in '。！？\?!':
                part = part + '。'

            sentences.append(part)

        # 按长度分组
        for sentence in sentences:
            if len(current) + len(sentence) <= max_length:
                current += sentence + " "
            else:
                if current.strip():
                    paragraphs.append(current.strip())
                current = sentence + " "

        if current.strip():
            paragraphs.append(current.strip())

        return paragraphs if paragraphs else [text]


class MiniMaxAPI:
    """MiniMax API 文本处理"""

    API_KEY = "sk-cp-XzaZl5R5kYt9bQ8Pess2o64FbD4wox5fRr6n4ZJqSWWL0Z5bMp5I2EIvRCHtJN_omeCiYHAUQjPN8FSenioRWcVoZ1hkbYcLz3gQ6JHmZa38GsqH6kLlnuQ"
    API_URL = "https://api.minimax.chat/v1/text/chatcompletion_pro"

    def __init__(self):
        self.headers = {
            "Authorization": f"Bearer {self.API_KEY}",
            "Content-Type": "application/json"
        }

    def polish_text(self, text, callback=None):
        """润色文本：添加标点、智能分段"""
        if not text or not text.strip():
            return text

        prompt = f"""请对以下转录文本进行润色处理：
1. 添加适当的标点符号（句号、逗号、引号等）
2. 智能分段（每段300-500字，根据语义完整性分段）
3. 修正明显的语音识别错误
4. 保持原始语义，不要过度修改

转录文本：
{text}

请直接输出润色后的文本，不需要任何解释。"""

        payload = {
            "model": "abab6.5s-chat",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 4096
        }

        try:
            response = requests.post(
                self.API_URL,
                headers=self.headers,
                json=payload,
                timeout=60
            )

            if response.status_code == 200:
                result = response.json()
                polished = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                if polished:
                    if callback:
                        callback(100)
                    return polished.strip()
            else:
                print(f"MiniMax API 错误: {response.status_code} - {response.text}")
                return text

        except Exception as e:
            print(f"MiniMax API 请求失败: {e}")
            return text

        return text

    def polish_segments(self, segments, callback=None):
        """分段润色（处理长文本）"""
        polished_segments = []
        total = len(segments)

        for i, segment in enumerate(segments):
            if callback:
                callback(int((i + 1) / total * 100))

            if segment.strip():
                # 每段单独润色
                polished = self.polish_text(segment)
                polished_segments.append(polished)
            else:
                polished_segments.append(segment)

        return polished_segments


class WordExporter:
    """Word 文档导出器"""

    def __init__(self):
        self.Document = None
        self._load_docx()

    def _load_docx(self):
        """延迟加载 python-docx"""
        try:
            from docx import Document
            self.Document = Document
            return True
        except ImportError:
            return False

    def create_document(self, paragraphs, output_path, title="转录文档"):
        """创建 Word 文档"""
        if self.Document is None:
            if not self._load_docx():
                return False, "请安装 python-docx: pip install python-docx"

        try:
            doc = self.Document()
            doc.core_properties.title = title

            # 添加标题
            heading = doc.add_heading(title, 0)
            heading.alignment = 1  # 居中

            # 添加时间戳
            doc.add_paragraph(f"创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            # 添加段落
            for i, para_text in enumerate(paragraphs):
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            # 保存
            doc.save(output_path)
            return True, f"文档已保存: {output_path}"

        except Exception as e:
            return False, f"保存失败: {str(e)}"

    def append_to_document(self, paragraphs, output_path):
        """追加内容到现有文档"""
        if self.Document is None:
            if not self._load_docx():
                return False, "请安装 python-docx: pip install python-docx"

        try:
            doc = self.Document(output_path)

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            doc.save(output_path)
            return True, f"已追加内容到: {output_path}"

        except Exception as e:
            return False, f"追加失败: {str(e)}"


class ProgressTracker:
    """进度追踪器"""

    def __init__(self, checkpoint_file=None):
        self.checkpoint_file = checkpoint_file
        self.state = {
            "video_path": "",
            "output_path": "",
            "config": {},
            "state": {
                "current_chunk": 0,
                "total_chunks": 0,
                "transcribed_text": [],
                "last_update": "",
                "stage": "idle"  # idle, extracting, transcribing, formatting, completed
            }
        }

    def save_checkpoint(self):
        """保存检查点"""
        if self.checkpoint_file:
            self.state["state"]["last_update"] = datetime.now().isoformat()
            try:
                with open(self.checkpoint_file, 'w', encoding='utf-8') as f:
                    json.dump(self.state, f, ensure_ascii=False, indent=2)
            except Exception as e:
                print(f"保存检查点失败: {e}")

    def load_checkpoint(self):
        """加载检查点"""
        if self.checkpoint_file and os.path.exists(self.checkpoint_file):
            try:
                with open(self.checkpoint_file, 'r', encoding='utf-8') as f:
                    self.state = json.load(f)
                return True
            except Exception:
                pass
        return False

    def can_resume(self, video_path):
        """检查是否可以恢复"""
        if not self.load_checkpoint():
            return False
        return (
            self.state["state"]["stage"] != "completed" and
            self.state["state"]["stage"] != "idle" and
            self.state.get("video_path") == video_path and
            self.state["state"]["current_chunk"] < self.state["state"]["total_chunks"]
        )


# ============================================================================
# 工作线程
# ============================================================================

class TranscribeWorker(QThread):
    """转录工作线程"""
    progress = pyqtSignal(str, int, int, str)  # stage, current, total, message
    log = pyqtSignal(str, str)  # level, message
    finished = pyqtSignal(bool, str)  # success, message
    memory_warning = pyqtSignal(int)  # memory percentage

    def __init__(self, video_path, output_path, config, checkpoint_file):
        super().__init__()
        self.video_path = video_path
        self.output_path = output_path
        self.config = config
        self.checkpoint_file = checkpoint_file
        self.is_paused = False
        self.is_stopped = False
        self.last_progress_time = time.time()
        self.stuck_threshold = 30  # 30秒无进度视为卡住

    def run(self):
        """执行转录任务"""
        try:
            self._run_transcription()
        except Exception as e:
            self.log.emit("ERROR", f"任务异常: {str(e)}\n{traceback.format_exc()}")
            self.finished.emit(False, f"任务异常: {str(e)}")

    def pause(self):
        self.is_paused = True

    def resume(self):
        self.is_paused = False

    def stop(self):
        self.is_stopped = True

    def _wait_if_paused(self):
        """暂停检查"""
        while self.is_paused and not self.is_stopped:
            QThread.sleep(0.5)

    def _check_stuck(self):
        """检查是否卡住"""
        elapsed = time.time() - self.last_progress_time
        if elapsed > self.stuck_threshold:
            self.log.emit("WARNING", f"检测到可能卡住（已处理 {elapsed:.0f} 秒无进度）")
            return True
        return False

    def _update_progress(self, stage, current, total, message=""):
        self.last_progress_time = time.time()
        self.progress.emit(stage, current, total, message)

        # 检查内存
        memory_percent = psutil.virtual_memory().percent
        if memory_percent > 80:
            self.memory_warning.emit(memory_percent)

    def _run_transcription(self):
        """执行转录流程"""
        tracker = ProgressTracker(self.checkpoint_file)

        # 检查断点续传
        can_resume = tracker.can_resume(self.video_path)
        if can_resume:
            self.log.emit("INFO", f"发现未完成的任务，将从第 {tracker.state['state']['current_chunk'] + 1} 个片段继续")
            choice = QMessageBox.question(
                None, "断点续传",
                "发现未完成的任务，是否继续？",
                QMessageBox.Yes | QMessageBox.No
            )
            if choice != QMessageBox.Yes:
                tracker.state["state"]["current_chunk"] = 0
                tracker.state["state"]["transcribed_text"] = []

        # 初始化组件
        audio_extractor = AudioExtractor()
        transcriber = Transcriber(
            model_size=self.config.get("model", "small"),
            device="cuda",
            compute_type=self.config.get("compute_type", "int8_float16")
        )
        punctuation_adder = PunctuationAdder()
        word_exporter = WordExporter()

        # 确保输出目录存在
        os.makedirs(os.path.dirname(self.output_path) or ".", exist_ok=True)

        # 获取音频分段配置
        chunk_duration = self.config.get("chunk_duration", 300)  # 默认5分钟

        # 阶段1: 音频提取（整视频）
        self._update_progress("extracting", 0, 100, "正在提取音频...")
        self.log.emit("INFO", "开始音频提取...")

        temp_audio = os.path.join(
            os.path.dirname(self.output_path),
            f".temp_audio_{os.path.basename(self.video_path)}.wav"
        )

        success, msg = audio_extractor.extract_audio(
            self.video_path, temp_audio,
            callback=lambda p: self._update_progress("extracting", p, 100, f"音频提取中: {p}%")
        )

        if not success:
            self.log.emit("ERROR", msg)
            self.finished.emit(False, msg)
            return

        self.log.emit("INFO", f"音频提取完成: {temp_audio}")

        # 获取音频总时长
        total_duration = audio_extractor.get_audio_duration(temp_audio)
        total_chunks = max(1, int(total_duration / chunk_duration) + 1)

        self.log.emit("INFO", f"音频总时长: {total_duration:.0f}秒，分为 {total_chunks} 个片段")

        # 更新检查点
        tracker.state["video_path"] = self.video_path
        tracker.state["output_path"] = self.output_path
        tracker.state["config"] = self.config
        tracker.state["state"]["total_chunks"] = total_chunks

        if not can_resume:
            tracker.state["state"]["current_chunk"] = 0
            tracker.state["state"]["transcribed_text"] = []

        tracker.save_checkpoint()

        # 阶段2: 分段转录
        all_texts = tracker.state["state"]["transcribed_text"].copy()
        start_chunk = tracker.state["state"]["current_chunk"]

        for chunk_idx in range(start_chunk, total_chunks):
            self._wait_if_paused()
            if self.is_stopped:
                self.log.emit("INFO", "用户停止任务")
                tracker.save_checkpoint()
                self.finished.emit(False, "任务已停止")
                return

            chunk_start = chunk_idx * chunk_duration
            chunk_end = min((chunk_idx + 1) * chunk_duration, total_duration)

            self._update_progress(
                "transcribing", chunk_idx + 1, total_chunks,
                f"正在转录第 {chunk_idx + 1}/{total_chunks} 个片段..."
            )
            self.log.emit("INFO", f"开始转录片段 {chunk_idx + 1}/{total_chunks} (时间: {chunk_start:.0f}s - {chunk_end:.0f}s)")

            success, msg, text = transcriber.transcribe(
                temp_audio,
                chunk_start=chunk_start,
                chunk_duration=chunk_duration
            )

            if not success:
                self.log.emit("ERROR", f"片段 {chunk_idx + 1} 转录失败: {msg}")
                # 继续处理下一个片段
                continue

            # 添加标点
            if self.config.get("add_punctuation", True):
                text = punctuation_adder.add_punctuation(text)

            all_texts.append(text)

            # 增量保存
            tracker.state["state"]["current_chunk"] = chunk_idx + 1
            tracker.state["state"]["transcribed_text"] = all_texts
            tracker.save_checkpoint()

            self.log.emit("INFO", f"片段 {chunk_idx + 1} 完成，文本长度: {len(text)} 字符")

        # 阶段3: 格式化并导出
        self._update_progress("formatting", 0, 100, "正在生成文档...")

        # 合并所有文本
        full_text = " ".join(all_texts)
        self.log.emit("INFO", f"原始文本长度: {len(full_text)} 字符")

        # MiniMax 润色（如果启用）
        use_polish = self.config.get("use_polish", False)
        if use_polish:
            self.log.emit("INFO", "正在使用 MiniMax 润色文本...")
            self._update_progress("formatting", 10, 100, "正在润色文本...")
            try:
                minimax = MiniMaxAPI()
                # 分段处理，避免超出 token 限制
                segments = punctuation_adder.split_paragraphs(full_text, max_length=800)
                polished_segments = minimax.polish_segments(
                    segments,
                    callback=lambda p: self._update_progress("formatting", 10 + int(p * 0.5), 100, f"润色进度: {p}%")
                )
                full_text = " ".join(polished_segments)
                self.log.emit("INFO", f"润色后文本长度: {len(full_text)} 字符")
            except Exception as e:
                self.log.emit("WARNING", f"MiniMax 润色失败: {e}，使用原始文本")
        else:
            # 本地标点处理
            full_text = punctuation_adder.add_punctuation(full_text)

        self._update_progress("formatting", 60, 100, "正在分段...")

        # 智能分段
        paragraphs = punctuation_adder.split_paragraphs(full_text, max_length=500)

        self._update_progress("formatting", 80, 100, "正在导出 Word...")

        # 如果有现有文档（断点续传），先删除
        if can_resume and os.path.exists(self.output_path):
            try:
                os.remove(self.output_path)
            except:
                pass

        # 导出 Word
        success, msg = word_exporter.create_document(
            paragraphs, self.output_path,
            title=f"视频转录 - {os.path.basename(self.video_path)}"
        )

        if success:
            self._update_progress("completed", 100, 100, "转录完成！")
            self.log.emit("INFO", msg)
            self.finished.emit(True, msg)

            # 清理临时文件
            try:
                os.remove(temp_audio)
            except:
                pass

            # 标记完成
            tracker.state["state"]["stage"] = "completed"
            tracker.save_checkpoint()
        else:
            self.log.emit("ERROR", msg)
            self.finished.emit(False, msg)


# ============================================================================
# 主窗口
# ============================================================================

class MainWindow(QMainWindow):
    """主窗口"""

    def __init__(self):
        super().__init__()
        self.worker = None
        self.is_processing = False
        self.memory_timer = None

        self.setWindowTitle("视频转文字工具")
        self.setMinimumSize(900, 700)

        # 样式表
        self.setStyleSheet("""
            QMainWindow {
                background-color: #1e1e1e;
                color: #d4d4d4;
            }
            QGroupBox {
                border: 1px solid #3c3c3c;
                border-radius: 5px;
                margin-top: 10px;
                font-weight: bold;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
            QLabel {
                color: #d4d4d4;
            }
            QLineEdit {
                background-color: #2d2d2d;
                border: 1px solid #3c3c3c;
                border-radius: 3px;
                padding: 5px;
                color: #d4d4d4;
            }
            QPushButton {
                background-color: #0e639c;
                border: none;
                border-radius: 3px;
                padding: 8px 16px;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1177bb;
            }
            QPushButton:pressed {
                background-color: #0d5a8f;
            }
            QPushButton:disabled {
                background-color: #3c3c3c;
                color: #808080;
            }
            QPushButton#btn_stop {
                background-color: #d73000;
            }
            QPushButton#btn_stop:hover {
                background-color: #f04000;
            }
            QPushButton#btn_stop:disabled {
                background-color: #5c3c3c;
            }
            QComboBox {
                background-color: #2d2d2d;
                border: 1px solid #3c3c3c;
                border-radius: 3px;
                padding: 5px;
                color: #d4d4d4;
            }
            QCheckBox {
                color: #d4d4d4;
            }
            QProgressBar {
                border: 1px solid #3c3c3c;
                border-radius: 3px;
                background-color: #2d2d2d;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #0e639c;
            }
            QTextEdit {
                background-color: #1e1e1e;
                border: 1px solid #3c3c3c;
                color: #d4d4d4;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 12px;
            }
            QStatusBar {
                background-color: #007acc;
                color: white;
            }
        """)

        self._init_ui()

        # 窗口关闭事件
        self.closeEvent = self.on_close_event

    def _init_ui(self):
        """初始化界面"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setSpacing(10)
        layout.setContentsMargins(15, 15, 15, 15)

        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QGridLayout()

        file_layout.addWidget(QLabel("视频文件:"), 0, 0)
        self.video_path_edit = QLineEdit()
        self.video_path_edit.setPlaceholderText("选择视频文件...")
        file_layout.addWidget(self.video_path_edit, 0, 1)
        self.btn_browse_video = QPushButton("浏览")
        self.btn_browse_video.clicked.connect(self.browse_video)
        file_layout.addWidget(self.btn_browse_video, 0, 2)

        file_layout.addWidget(QLabel("输出位置:"), 1, 0)
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setPlaceholderText("选择输出目录...")
        file_layout.addWidget(self.output_path_edit, 1, 1)
        self.btn_browse_output = QPushButton("浏览")
        self.btn_browse_output.clicked.connect(self.browse_output)
        file_layout.addWidget(self.btn_browse_output, 1, 2)

        file_group.setLayout(file_layout)
        layout.addWidget(file_group)

        # 参数设置区域
        config_group = QGroupBox("参数设置")
        config_layout = QGridLayout()

        config_layout.addWidget(QLabel("Whisper 模型:"), 0, 0)
        self.model_combo = QComboBox()
        self.model_combo.addItems(["tiny", "small", "medium"])
        self.model_combo.setCurrentText("small")
        config_layout.addWidget(self.model_combo, 0, 1)

        config_layout.addWidget(QLabel("量化类型:"), 0, 2)
        self.compute_type_combo = QComboBox()
        self.compute_type_combo.addItems([
            "int8_float16 (推荐-省内存)",
            "int8 (省内存-速度慢)",
            "float16 (高质量-显存高)"
        ])
        self.compute_type_combo.setCurrentIndex(0)
        config_layout.addWidget(self.compute_type_combo, 0, 3)

        config_layout.addWidget(QLabel("分段时长(分钟):"), 1, 0)
        self.chunk_duration_combo = QComboBox()
        self.chunk_duration_combo.addItems(["5", "10", "15"])
        self.chunk_duration_combo.setCurrentText("5")
        config_layout.addWidget(self.chunk_duration_combo, 1, 1)

        self.punctuation_check = QCheckBox("本地标点")
        self.punctuation_check.setChecked(True)
        config_layout.addWidget(self.punctuation_check, 1, 2)

        self.polish_check = QCheckBox("MiniMax 润色 (需联网)")
        self.polish_check.setChecked(False)
        config_layout.addWidget(self.polish_check, 1, 3)

        config_group.setLayout(config_layout)
        layout.addWidget(config_group)

        # 操作按钮区域
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        self.btn_start = QPushButton("开始转录")
        self.btn_start.clicked.connect(self.start_transcribe)
        btn_layout.addWidget(self.btn_start)

        self.btn_pause = QPushButton("暂停")
        self.btn_pause.clicked.connect(self.toggle_pause)
        self.btn_pause.setEnabled(False)
        btn_layout.addWidget(self.btn_pause)

        self.btn_stop = QPushButton("停止")
        self.btn_stop.setObjectName("btn_stop")
        self.btn_stop.clicked.connect(self.stop_transcribe)
        self.btn_stop.setEnabled(False)
        btn_layout.addWidget(self.btn_stop)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # 进度显示区域
        progress_group = QGroupBox("处理进度")
        progress_layout = QVBoxLayout()

        self.stage_label = QLabel("状态: 待机")
        progress_layout.addWidget(self.stage_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar)

        self.chunk_label = QLabel("片段: 0/0")
        progress_layout.addWidget(self.chunk_label)

        self.time_label = QLabel("预估剩余: --")
        progress_layout.addWidget(self.time_label)

        self.memory_label = QLabel("内存使用: --")
        progress_layout.addWidget(self.memory_label)

        progress_group.setLayout(progress_layout)
        layout.addWidget(progress_group)

        # 日志区域
        log_group = QGroupBox("日志")
        log_layout = QVBoxLayout()

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        log_layout.addWidget(self.log_text)

        log_group.setLayout(log_layout)
        layout.addWidget(log_group)

        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("就绪 | GPU: 检测中...")

        # 内存监控定时器
        self.memory_timer = QTimer()
        self.memory_timer.timeout.connect(self.update_memory_display)
        self.memory_timer.start(5000)

        self.log_message("INFO", "视频转文字工具已启动")
        self.log_message("INFO", "请选择视频文件开始转录")

    def browse_video(self):
        """选择视频文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择视频文件",
            "",
            "视频文件 (*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.ts *.m2ts);;所有文件 (*.*)"
        )
        if file_path:
            self.video_path_edit.setText(file_path)

            # 自动设置输出路径
            if not self.output_path_edit.text():
                default_output = os.path.join(
                    os.path.dirname(file_path),
                    os.path.splitext(os.path.basename(file_path))[0] + ".docx"
                )
                self.output_path_edit.setText(default_output)

    def browse_output(self):
        """选择输出目录"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存文档",
            self.output_path_edit.text() or "output.docx",
            "Word 文档 (*.docx)"
        )
        if file_path:
            self.output_path_edit.setText(file_path)

    def start_transcribe(self):
        """开始转录"""
        video_path = self.video_path_edit.text().strip()
        output_path = self.output_path_edit.text().strip()

        if not video_path:
            QMessageBox.warning(self, "提示", "请选择视频文件")
            return

        if not output_path:
            QMessageBox.warning(self, "提示", "请选择输出位置")
            return

        if not os.path.exists(video_path):
            QMessageBox.warning(self, "警告", "视频文件不存在")
            return

        # 获取配置
        compute_type_map = {
            0: "int8_float16",
            1: "int8",
            2: "float16"
        }
        config = {
            "model": self.model_combo.currentText(),
            "compute_type": compute_type_map.get(self.compute_type_combo.currentIndex(), "int8_float16"),
            "chunk_duration": int(self.chunk_duration_combo.currentText()) * 60,
            "add_punctuation": self.punctuation_check.isChecked(),
            "use_polish": self.polish_check.isChecked()
        }

        # 检查点文件
        checkpoint_file = os.path.join(
            os.path.dirname(output_path),
            f".checkpoint_{os.path.basename(output_path)}.json"
        )

        # 创建工作线程
        self.worker = TranscribeWorker(video_path, output_path, config, checkpoint_file)
        self.worker.progress.connect(self.on_progress)
        self.worker.log.connect(self.log_message)
        self.worker.finished.connect(self.on_finished)
        self.worker.memory_warning.connect(self.on_memory_warning)
        self.worker.start()

        # 更新UI状态
        self.is_processing = True
        self.btn_start.setEnabled(False)
        self.btn_pause.setEnabled(True)
        self.btn_stop.setEnabled(True)
        self.status_bar.showMessage("处理中...")

        self.log_message("INFO", f"开始处理: {video_path}")

    def toggle_pause(self):
        """暂停/继续"""
        if self.worker:
            if self.btn_pause.text() == "暂停":
                self.worker.pause()
                self.btn_pause.setText("继续")
                self.status_bar.showMessage("已暂停")
            else:
                self.worker.resume()
                self.btn_pause.setText("暂停")
                self.status_bar.showMessage("处理中...")

    def stop_transcribe(self):
        """停止转录"""
        if self.worker:
            choice = QMessageBox.question(
                self, "确认停止",
                "确定要停止转录吗？\n进度已保存，下次可继续。",
                QMessageBox.Yes | QMessageBox.No
            )
            if choice == QMessageBox.Yes:
                self.worker.stop()
                self.log_message("INFO", "正在停止任务...")

    def on_progress(self, stage, current, total, message):
        """进度更新"""
        self.stage_label.setText(f"状态: {message}")

        if stage == "extracting":
            self.progress_bar.setValue(current)
        elif stage == "transcribing":
            if total > 0:
                progress = int(current / total * 100)
                self.progress_bar.setValue(progress)
            self.chunk_label.setText(f"片段: {current}/{total}")
        elif stage == "formatting":
            self.progress_bar.setValue(current)
        elif stage == "completed":
            self.progress_bar.setValue(100)

    def log_message(self, level, message):
        """添加日志消息"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        color_map = {
            "INFO": "#4ec9b0",
            "WARNING": "#dcdcaa",
            "ERROR": "#f14c4c"
        }
        color = color_map.get(level, "#d4d4d4")

        self.log_text.append(
            f'<span style="color: #808080;">[{timestamp}]</span> '
            f'<span style="color: {color};">{level}: {message}</span>'
        )

        # 自动滚动
        cursor = self.log_text.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.log_text.setTextCursor(cursor)

    def on_finished(self, success, message):
        """任务完成"""
        self.is_processing = False
        self.btn_start.setEnabled(True)
        self.btn_pause.setEnabled(False)
        self.btn_pause.setText("暂停")
        self.btn_stop.setEnabled(False)

        if success:
            self.status_bar.showMessage("转录完成")
            self.log_message("INFO", f"✓ {message}")
            QMessageBox.information(self, "完成", message)
        else:
            self.status_bar.showMessage("任务异常")
            if "停止" not in message:
                self.log_message("ERROR", f"✗ {message}")
                QMessageBox.warning(self, "异常", message)

    def on_memory_warning(self, percent):
        """内存警告"""
        self.memory_label.setText(f"<span style='color: #f14c4c;'>内存使用: {percent}% (较高)</span>")
        if percent > 90:
            self.log_message("WARNING", f"内存使用率过高 ({percent}%)，建议关闭其他程序")

    def update_memory_display(self):
        """更新内存显示"""
        memory = psutil.virtual_memory()
        self.memory_label.setText(f"内存使用: {memory.percent}% ({memory.used // (1024**3)}/{memory.total // (1024**3)}GB)")

    def on_close_event(self, event):
        """窗口关闭事件"""
        if self.is_processing:
            choice = QMessageBox.question(
                self, "确认退出",
                "正在处理中，确定要退出吗？\n进度已保存，可稍后继续。",
                QMessageBox.Yes | QMessageBox.No
            )
            if choice == QMessageBox.Yes:
                if self.worker:
                    self.worker.stop()
                    self.worker.wait(3000)
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()


# ============================================================================
# 程序入口
# ============================================================================

def main():
    app = QApplication(sys.argv)
    app.setApplicationName("视频转文字工具")

    window = MainWindow()
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
