#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
视频转文字工具 - Web 服务器入口
"""
import os
import sys
import json
import time
import uuid
import queue
import threading
import traceback
import subprocess
import logging

# 设置工作目录
script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)
os.chdir(script_dir)  # 切换到脚本目录

# 导入核心模块
from core import (
    AudioExtractor, Transcriber, PunctuationAdder,
    MiniMaxAPI, WordExporter, ProgressManager
)

# 启动时立即初始化日志
LOG_FILE = os.path.join(script_dir, 'web_server.log')
try:
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(LOG_FILE, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
except Exception as e:
    with open(LOG_FILE, 'w', encoding='utf-8') as f:
        f.write(f"[ERROR] 日志初始化失败: {e}\n")

logger = logging.getLogger(__name__)

# 启动时自动清理旧日志
def cleanup_old_logs():
    """启动时清理旧日志"""
    try:
        task_log = os.path.join(script_dir, "输出文件", "task.log")
        if os.path.exists(task_log):
            with open(task_log, 'w', encoding='utf-8') as f:
                f.write("")
        
        progress_file = os.path.join(script_dir, '输出文件', 'task_progress.json')
        if os.path.exists(progress_file):
            os.remove(progress_file)
        
        logger.info("旧日志已清理")
    except Exception as e:
        logger.warning(f"清理日志失败: {e}")

cleanup_old_logs()

logger.info("程序启动成功")

from flask import Flask, request, send_file, jsonify, render_template, Response
import psutil

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# 创建 Flask 应用
app = Flask(__name__, static_folder='web', template_folder='web')
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024 * 1024  # 10GB max
app.config['UPLOAD_FOLDER'] = '输出文件'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 全局状态
class TaskState:
    def __init__(self):
        self.is_running = False
        self.is_paused = False
        self.is_stopped = False
        self.progress = 0
        self.stage = ""
        self.current_chunk = 0
        self.total_chunks = 0
        self.time_left = ""
        self.last_update = time.time()
        self.current_task_id = None

state = TaskState()

# SSE 消息队列
message_queue = queue.Queue()

def send_sse(data, task_id=None):
    """发送 Server-Sent Events"""
    try:
        if task_id:
            data['task_id'] = task_id
        message_queue.put_nowait(json.dumps(data))
    except:
        pass
    
    if data.get("type") == "log":
        level = data.get("level", "info")
        msg = data.get("message", "")
        log_file = os.path.join(script_dir, "输出文件", "task.log")
        try:
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(f"[{level.upper()}] {msg}\n")
        except:
            pass

def sse_generator():
    """SSE 流生成器"""
    try:
        while True:
            try:
                data = message_queue.get(timeout=5)
                yield f"data: {data}\n\n"
            except queue.Empty:
                yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"
    except GeneratorExit:
        pass


def run_transcription(video_path, output_dir, config, task_id=None):
    """执行转录任务"""
    global state
    task_start_time = time.time()

    # 清空任务日志
    log_file = os.path.join(script_dir, "输出文件", "task.log")
    try:
        with open(log_file, "w", encoding="utf-8") as f:
            f.write("")
    except:
        pass

    state.is_running = True
    state.is_paused = False
    state.is_stopped = False
    state.progress = 0
    state.stage = "初始化"
    state.current_chunk = 0
    state.total_chunks = 0
    state.current_task_id = task_id

    def send_task_sse(data):
        send_sse(data, task_id)
        if data.get("type") in ["progress", "log", "complete", "error"]:
            save_task_progress(task_id, data)

    temp_files = []
    progress_file = None

    try:
        send_task_sse({"type": "log", "level": "info", "message": f"========== 任务开始 =========="})
        send_task_sse({"type": "log", "level": "info", "message": f"视频路径: {video_path}"})
        send_task_sse({"type": "log", "level": "info", "message": f"视频大小: {os.path.getsize(video_path) / 1024 / 1024:.2f} MB"})
        send_task_sse({"type": "log", "level": "info", "message": f"Whisper模型: {config['model']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"分段时长: {config['chunk_duration']} 分钟"})
        send_task_sse({"type": "log", "level": "info", "message": f"启用标点: {config['punctuation']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"启用润色: {config['polish']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"输出目录: {output_dir}"})
        send_task_sse({"type": "log", "level": "info", "message": "=" * 40})

        # 初始化组件
        send_task_sse({"type": "log", "level": "info", "message": "开始初始化组件..."})
        audio_extractor = AudioExtractor()
        transcriber = Transcriber(config["model"])
        punctuation_adder = PunctuationAdder()
        word_exporter = WordExporter()

        output_name = config.get("output_name", "转录文档")
        output_path = os.path.join(output_dir, output_name + ".docx")
        progress_file = os.path.join(output_dir, f".{output_name}_progress.json")

        # 阶段1: 音频提取
        send_task_sse({"type": "log", "level": "info", "message": "正在提取音频..."})
        state.stage = "音频提取"

        temp_audio = os.path.join(output_dir, f".temp_{uuid.uuid4().hex}.wav")
        temp_files.append(temp_audio)

        success, msg = audio_extractor.extract_audio(
            video_path, temp_audio,
            progress_callback=lambda p: send_task_sse({
                "type": "progress",
                "progress": int(p * 0.05),
                "stage": f"音频提取 {p}%"
            })
        )

        if not success:
            raise Exception(f"音频提取失败: {msg}")

        send_task_sse({"type": "log", "level": "info", "message": "音频提取完成"})

        # 加载模型
        send_task_sse({"type": "log", "level": "info", "message": "加载 Whisper 模型..."})
        success, msg = transcriber.load_model()
        if not success:
            raise Exception(msg)

        send_task_sse({"type": "log", "level": "info", "message": f"模型加载完成: {transcriber.model_size}"})

        # 获取总时长
        total_duration = audio_extractor._get_duration(temp_audio)
        chunk_duration = int(config.get("chunk_duration", 5)) * 60
        total_chunks = max(1, int(total_duration / chunk_duration) + (1 if total_duration % chunk_duration > 0 else 0))
        state.total_chunks = total_chunks

        send_task_sse({"type": "log", "level": "info", "message": f"音频时长: {total_duration:.0f}秒, 分为 {total_chunks} 个片段"})

        # 检查断点续录
        start_chunk_idx = 0
        completed_texts = []
        if os.path.exists(progress_file):
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    saved = json.load(f)
                    saved_video = saved.get('video_path', '')
                    
                    if not os.path.isabs(saved_video):
                        saved_video_abs = os.path.join(script_dir, saved_video)
                    else:
                        saved_video_abs = saved_video
                    
                    video_exists = os.path.exists(saved_video_abs)
                    
                    original_name = os.path.basename(video_path)
                    saved_original_name = os.path.basename(saved_video)
                    if '_' in original_name and '_' in saved_original_name:
                        orig_short = '_'.join(original_name.split('_')[1:])
                        saved_short = '_'.join(saved_original_name.split('_')[1:])
                        name_match = (orig_short == saved_short)
                    else:
                        name_match = (original_name == saved_original_name)
                    
                    if video_exists and name_match and saved.get('total_chunks') == total_chunks:
                        start_chunk_idx = saved.get('last_completed_chunk', 0) + 1
                        completed_texts = saved.get('texts', [])
                        completed_count = len(completed_texts)
                        resume_progress = 5 + int(completed_count / total_chunks * 60)
                        send_task_sse({
                            "type": "progress",
                            "progress": resume_progress,
                            "stage": f"续传片段 {start_chunk_idx + 1}/{total_chunks}",
                            "current": completed_count,
                            "total": total_chunks,
                            "time_left": ""
                        })
                        send_task_sse({"type": "log", "level": "info", "message": f"检测到进度文件，从片段 {start_chunk_idx + 1} 继续（已完成 {completed_count} 个片段）"})
            except:
                pass

        # 阶段2: 分段转录
        start_time = time.time()
        base_progress = 5

        for chunk_idx in range(start_chunk_idx, total_chunks):
            if state.is_stopped:
                send_task_sse({"type": "log", "level": "warning", "message": "任务已停止"})
                save_progress(progress_file, {
                    'video_path': video_path,
                    'total_chunks': total_chunks,
                    'last_completed_chunk': chunk_idx - 1,
                    'texts': completed_texts,
                    'polished': [],
                    'polish_completed_segments': 0,
                    'outlines': []
                })
                break

            while state.is_paused and not state.is_stopped:
                time.sleep(0.5)

            chunk_start = chunk_idx * chunk_duration
            chunk_end = min(chunk_start + chunk_duration, total_duration)
            actual_duration = chunk_end - chunk_start
            state.current_chunk = chunk_idx + 1
            state.stage = f"转录片段 {chunk_idx + 1}/{total_chunks}"

            segment_file = os.path.join(output_dir, f".segment_{chunk_idx:03d}.wav")
            temp_files.append(segment_file)

            send_task_sse({"type": "log", "level": "info", "message": f"切割片段 {chunk_idx + 1}/{total_chunks}..."})
            success, msg = audio_extractor.cut_segment(temp_audio, segment_file, chunk_start, actual_duration)
            if not success:
                send_task_sse({"type": "log", "level": "warning", "message": f"片段 {chunk_idx + 1} 切割失败: {msg}，跳过"})
                continue

            chunk_progress = (chunk_idx - start_chunk_idx + 1) / (total_chunks - start_chunk_idx) * 60
            elapsed = time.time() - start_time
            remaining = int(elapsed / (chunk_idx - start_chunk_idx + 1) * (total_chunks - chunk_idx - 1))
            time_left = f"{remaining // 60}分{remaining % 60}秒" if remaining >= 60 else f"{remaining}秒"
            
            send_task_sse({
                "type": "progress",
                "progress": base_progress + int(chunk_progress),
                "stage": state.stage,
                "current": chunk_idx + 1,
                "total": total_chunks,
                "time_left": time_left
            })

            send_task_sse({"type": "log", "level": "info", "message": f"开始转录片段 {chunk_idx + 1}/{total_chunks} ({chunk_start:.0f}-{chunk_end:.0f}秒)"})
            success, msg, text = transcriber.transcribe(segment_file, chunk_start=0, chunk_end=None)

            if success and text:
                paragraphs = punctuation_adder.split_paragraphs(text, 500)

                if chunk_idx == start_chunk_idx and not completed_texts:
                    success, msg = word_exporter.create_initial_document(output_path, title=f"视频转录 - {output_name}")
                    if not success:
                        send_task_sse({"type": "log", "level": "warning", "message": f"文档创建失败: {msg}"})

                success, msg = word_exporter.append_paragraphs(paragraphs, output_path)
                if success:
                    send_task_sse({"type": "log", "level": "info", "message": f"片段 {chunk_idx + 1} 已保存 ({len(text)} 字符)"})

                completed_texts.append(text)

                save_progress(progress_file, {
                    'video_path': video_path,
                    'total_chunks': total_chunks,
                    'last_completed_chunk': chunk_idx,
                    'texts': completed_texts,
                    'polished': [],
                    'polish_completed_segments': 0,
                    'outlines': []
                })
            else:
                send_task_sse({"type": "log", "level": "warning", "message": f"片段 {chunk_idx + 1} 转录失败: {msg}"})

            try:
                if os.path.exists(segment_file):
                    os.remove(segment_file)
                    temp_files.remove(segment_file)
            except:
                pass

        send_task_sse({"type": "log", "level": "info", "message": f"转录完成，共 {len(completed_texts)} 个片段"})

        if not completed_texts:
            raise Exception("没有成功转录任何内容")

        # 阶段3: 润色
        full_text = " ".join(completed_texts)
        send_task_sse({"type": "log", "level": "info", "message": f"原始文本: {len(full_text)} 字符"})
        state.stage = "后处理"
        video_outline = ""

        saved_polished = []
        saved_outlines = []
        polish_completed = 0
        if os.path.exists(progress_file):
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    saved_progress = json.load(f)
                    saved_polished = saved_progress.get('polished', [])
                    saved_outlines = saved_progress.get('outlines', [])
                    polish_completed = saved_progress.get('polish_completed_segments', 0)
            except:
                pass

        if config.get("polish", False):
            api_key = config.get("api_key", "").strip()
            if not api_key:
                send_task_sse({"type": "log", "level": "warning", "message": "未提供 MiniMax API Key，无法润色"})
                full_text = punctuation_adder.add_punctuation(full_text)
            else:
                send_task_sse({"type": "log", "level": "info", "message": "正在使用 MiniMax 润色..."})
                try:
                    minimax = MiniMaxAPI(api_key, model=config.get("minimax_model", "MiniMax-M2.5"))
                    segments = punctuation_adder.split_paragraphs(full_text, 800)
                    total_polish = len(segments)

                    if saved_polished and len(saved_polished) >= polish_completed:
                        polished = saved_polished[:polish_completed]
                        outlines = saved_outlines[:polish_completed]
                        start_polish_idx = polish_completed
                        send_task_sse({"type": "log", "level": "info", "message": f"检测到润色进度，从第 {polish_completed + 1} 段继续（共 {total_polish} 段）"})
                    else:
                        polished = []
                        outlines = []
                        start_polish_idx = 0

                    polish_start = time.time()
                    if start_polish_idx > 0:
                        elapsed = start_polish_idx * 30
                        polish_start = time.time() - elapsed

                    def polish_log(msg, level="info"):
                        send_task_sse({"type": "log", "level": level, "message": msg})

                    for i, seg in enumerate(segments):
                        if i < start_polish_idx:
                            continue

                        result, outline = minimax.polish_text(seg, log_callback=polish_log)
                        polished.append(result)
                        if outline:
                            outlines.append(outline)

                        current_completed = i + 1
                        save_progress(progress_file, {
                            'video_path': video_path,
                            'total_chunks': total_chunks,
                            'last_completed_chunk': total_chunks - 1,
                            'texts': completed_texts,
                            'polished': polished,
                            'polish_completed_segments': current_completed,
                            'outlines': outlines
                        })

                        elapsed = time.time() - polish_start
                        remaining = int(elapsed / current_completed * (total_polish - current_completed))
                        time_str = f"{remaining}秒" if remaining < 60 else f"{remaining // 60}分{remaining % 60}秒"
                        send_task_sse({
                            "type": "progress",
                            "progress": 60 + int(current_completed / total_polish * 20),
                            "stage": f"润色中 {current_completed}/{total_polish}",
                            "time_left": time_str
                        })

                    full_text = " ".join(polished)
                    video_outline = "\n".join(outlines) if outlines else ""
                    send_task_sse({"type": "log", "level": "info", "message": f"润色完成: {len(full_text)} 字符"})

                    has_punct = any(c in '。！？.?!' for c in full_text)
                    if not has_punct:
                        send_task_sse({"type": "log", "level": "info", "message": "润色内容无标点，正在添加标点..."})
                        full_text = punctuation_adder.add_punctuation(full_text)

                    send_task_sse({"type": "progress", "progress": 80, "stage": "后处理", "time_left": "即将完成"})
                except Exception as e:
                    send_task_sse({"type": "log", "level": "warning", "message": f"润色失败: {e}，使用原始文本"})
                    full_text = punctuation_adder.add_punctuation(full_text)
        else:
            full_text = punctuation_adder.add_punctuation(full_text)

        # 生成最终文档
        send_task_sse({"type": "log", "level": "info", "message": "正在生成最终文档..."})
        try:
            from docx import Document as DocXDocument
            
            # 创建新文档
            if os.path.exists(output_path):
                os.remove(output_path)
            success, msg = word_exporter.create_initial_document(output_path, title=f"视频转录 - {output_name}")
            if not success:
                raise Exception(msg)
            
            doc = DocXDocument(output_path)
            
            # 添加视频大纲（如果有）
            if video_outline and video_outline.strip():
                doc.add_paragraph()
                doc.add_heading("📋 视频大纲", level=2)
                for line in video_outline.split('\n'):
                    line = line.strip()
                    if line:
                        # 去掉可能存在的数字前缀（如 "1. " 或 "1、"）
                        import re
                        line = re.sub(r'^[\d]+\.\s*', '', line)
                        line = re.sub(r'^[\d]+、\s*', '', line)
                        p = doc.add_paragraph(line)
                        p.style = 'List Number'
                doc.add_paragraph()
            
            # 添加正文
            doc.add_heading("📝 正文内容", level=2)
            paragraphs = punctuation_adder.split_paragraphs(full_text, 500)
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.line_spacing = 1.5
            
            doc.save(output_path)
            if success:
                send_task_sse({"type": "log", "level": "info", "message": f"文档生成完成: {os.path.getsize(output_path)/1024:.1f} KB"})
        except Exception as e:
            send_task_sse({"type": "log", "level": "warning", "message": f"文档生成失败: {e}"})

        send_task_sse({"type": "progress", "progress": 100, "stage": "完成", "time_left": "0秒"})
        send_task_sse({"type": "complete", "output_path": output_path, "output_name": output_name, "memory": f"{psutil.virtual_memory().percent:.0f}%"})
        state.progress = 100
        state.stage = "完成"
        total_time = time.time() - task_start_time
        send_task_sse({"type": "log", "level": "info", "message": f"任务完成！总耗时: {total_time:.1f}秒 ({total_time/60:.1f}分钟)"})
        send_task_sse({"type": "log", "level": "info", "message": f"输出文件: {output_path}"})

    except Exception as e:
        send_task_sse({"type": "error", "message": str(e)})
        send_task_sse({"type": "log", "level": "error", "message": f"错误: {traceback.format_exc()}"})

    finally:
        for f in temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except:
                pass
        state.is_running = False


def save_progress(progress_file, data):
    """保存进度"""
    try:
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)
    except:
        pass


# 任务进度存储
task_progress_store = {}

def save_task_progress(task_id, data):
    task_progress_store[task_id] = data
    progress_file = os.path.join(script_dir, '输出文件', 'task_progress.json')
    try:
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(task_progress_store, f, ensure_ascii=False)
    except:
        pass

def load_task_progress(task_id):
    global task_progress_store
    if task_id in task_progress_store:
        return task_progress_store[task_id]
    progress_file = os.path.join(script_dir, '输出文件', 'task_progress.json')
    if os.path.exists(progress_file):
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                task_progress_store = json.load(f)
                return task_progress_store.get(task_id)
        except:
            pass
    return None


# ============================================================================
# Flask 路由
# ============================================================================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/transcribe', methods=['POST'])
def transcribe():
    if state.is_running:
        return jsonify({"error": "已有任务在运行"}), 400

    video = request.files.get('video')
    if not video:
        return jsonify({"error": "未上传视频文件"}), 400

    video_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex}_{video.filename}")
    video.save(video_path)

    video_dir = os.path.dirname(video_path)
    task_id = request.form.get('task_id', str(uuid.uuid4().hex[:8]))

    config = {
        "output_name": request.form.get('output_name', '转录文档'),
        "model": request.form.get('model', 'small'),
        "chunk_duration": int(request.form.get('chunk_duration', 5)),
        "punctuation": request.form.get('punctuation', 'true').lower() == 'true',
        "polish": request.form.get('polish', 'false').lower() == 'true',
        "api_key": request.form.get('api_key', ''),
        "minimax_model": request.form.get('minimax_model', 'MiniMax-M2.5'),
        "video_path": video_path,
        "output_dir": video_dir
    }

    thread = threading.Thread(target=run_transcription, args=(video_path, app.config['UPLOAD_FOLDER'], config, task_id))
    thread.daemon = False
    thread.start()

    return jsonify({"status": "started", "video_path": video_path, "task_id": task_id})

@app.route('/api/resume', methods=['POST'])
def resume():
    if state.is_running:
        return jsonify({"error": "已有任务在运行"}), 400

    video_path = request.form.get('video_path')
    output_name = request.form.get('output_name', '转录文档')
    if not video_path or not os.path.exists(video_path):
        return jsonify({"error": "视频文件不存在"}), 400

    output_dir = os.path.dirname(video_path)
    progress_file = os.path.join(output_dir, f".{output_name}_progress.json")
    if not os.path.exists(progress_file):
        return jsonify({"error": "没有可恢复的进度"}), 400

    try:
        with open(progress_file, 'r', encoding='utf-8') as f:
            saved = json.load(f)
    except:
        return jsonify({"error": "进度文件读取失败"}), 400

    if saved.get('video_path') != video_path:
        return jsonify({"error": "视频文件不匹配，无法恢复"}), 400

    task_id = request.form.get('task_id', str(uuid.uuid4().hex[:8]))

    config = {
        "output_name": output_name,
        "model": request.form.get('model', 'small'),
        "chunk_duration": int(request.form.get('chunk_duration', 5)),
        "punctuation": request.form.get('punctuation', 'true').lower() == 'true',
        "polish": request.form.get('polish', 'false').lower() == 'true',
        "api_key": request.form.get('api_key', ''),
        "minimax_model": request.form.get('minimax_model', 'MiniMax-M2.5'),
        "video_path": video_path,
        "output_dir": output_dir,
        "resume_from": saved.get('last_completed_chunk', -1) + 1
    }

    thread = threading.Thread(target=run_transcription, args=(video_path, output_dir, config, task_id))
    thread.daemon = False
    thread.start()

    return jsonify({"status": "resumed", "start_chunk": config["resume_from"], "task_id": task_id})

@app.route('/api/progress')
def progress():
    return Response(sse_generator(), mimetype='text/event-stream')

@app.route('/api/pause', methods=['POST'])
def pause():
    state.is_paused = not state.is_paused
    return jsonify({"paused": state.is_paused})

@app.route('/api/stop', methods=['POST'])
def stop():
    state.is_stopped = True
    return jsonify({"stopped": True})

@app.route('/api/status')
def status():
    return jsonify({
        "running": state.is_running,
        "paused": state.is_paused,
        "memory": f"{psutil.virtual_memory().percent:.0f}"
    })

@app.route('/api/poll_progress/<task_id>')
def poll_progress(task_id):
    progress = load_task_progress(task_id)
    if progress:
        return jsonify(progress)
    return jsonify({"type": "no_data"})

@app.route('/api/clear_logs', methods=['POST'])
def clear_logs():
    try:
        task_log = os.path.join(script_dir, "输出文件", "task.log")
        if os.path.exists(task_log):
            with open(task_log, 'w', encoding='utf-8') as f:
                f.write("")
        
        progress_file = os.path.join(script_dir, '输出文件', 'task_progress.json')
        if os.path.exists(progress_file):
            os.remove(progress_file)
        
        global task_progress_store
        task_progress_store = {}
        
        return jsonify({"success": True, "message": "日志已清理"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/download/<path:filename>')
def download(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(file_path):
        file_path = file_path + '.docx'
    return send_file(file_path, as_attachment=True)

@app.route('/api/shutdown', methods=['POST'])
def shutdown():
    """关闭服务器"""
    import threading
    def do_shutdown():
        time.sleep(0.5)  # 等待响应发送
        os._exit(0)
    threading.Thread(target=do_shutdown, daemon=True).start()
    return jsonify({"success": True, "message": "服务器正在关闭..."})


@app.route('/api/minimax/test', methods=['POST'])
def minimax_test():
    try:
        data = request.get_json()
        api_key = data.get('api_key', '').strip()

        if not api_key:
            return jsonify({"success": False, "error": "API Key 不能为空"})

        import urllib.request

        url = "https://api.minimaxi.com/anthropic/v1/messages"
        headers = {
            "x-api-key": api_key,
            "Content-Type": "application/json",
            "anthropic-version": "2023-06-01"
        }
        payload = json.dumps({
            "model": "MiniMax-M2.5",
            "max_tokens": 10,
            "messages": [{"role": "user", "content": "hi"}]
        }).encode('utf-8')

        req = urllib.request.Request(url, data=payload, headers=headers)
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                result = json.loads(response.read().decode('utf-8'))
                return jsonify({
                    "success": True,
                    "models": [
                        {"id": "MiniMax-M2.7"},
                        {"id": "MiniMax-M2.5"},
                        {"id": "MiniMax-M2.1"}
                    ]
                })
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8') if e.fp else ''
            try:
                error_json = json.loads(error_body)
                error_msg = error_json.get('error', {}).get('message', error_body) if isinstance(error_json.get('error'), dict) else error_json.get('error', error_body)
            except:
                error_msg = error_body if error_body else f"HTTP {e.code}"
            return jsonify({"success": False, "error": error_msg})
        except Exception as e:
            return jsonify({"success": False, "error": str(e)})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# ============================================================================
# 启动
# ============================================================================

if __name__ == '__main__':
    import webbrowser
    import logging

    log = logging.getLogger('werkzeug')
    log.setLevel(logging.ERROR)
    app.logger.setLevel(logging.ERROR)

    print("=" * 50)
    print("视频转文字工具 - Web 版")
    print("=" * 50)
    print("正在启动服务器...")
    print("请访问 http://127.0.0.1:5000")
    print("按 Ctrl+C 可关闭服务器")
    print("=" * 50)

    def open_browser():
        # 等待服务器完全启动后再打开浏览器
        import socket
        for _ in range(30):  # 最多等待30秒
            try:
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                sock.settimeout(1)
                result = sock.connect_ex(('127.0.0.1', 5000))
                sock.close()
                if result == 0:
                    break
            except:
                pass
            time.sleep(1)
        webbrowser.open('http://localhost:5000')

    threading.Thread(target=open_browser, daemon=True).start()

    # 无限循环启动，除非收到 KeyboardInterrupt
    restart_count = 0
    while True:
        try:
            restart_count += 1
            if restart_count > 1:
                print(f"\n[重启次数: {restart_count - 1}]")
                print("重新启动服务器...")
            app.run(host='0.0.0.0', port=5000, debug=False, threaded=True, use_reloader=False)
            # 如果正常退出（没有异常），跳出循环
            break
        except KeyboardInterrupt:
            print("\n程序已停止")
            break
        except GeneratorExit:
            print("\n客户端断开连接，继续运行...")
            continue
        except OSError as e:
            # 端口被占用等系统错误
            print(f"\n系统错误: {e}")
            if "10048" in str(e) or "Address already in use" in str(e):
                print("端口被占用，尝试关闭后重启...")
                import socket
                try:
                    s = socket.socket()
                    s.bind(('0.0.0.0', 5000))
                    s.close()
                except:
                    pass
            print("5秒后重新启动...")
            time.sleep(5)
            continue
        except BaseException as e:
            # 捕获所有其他异常，打印日志但不退出
            print(f"\n程序异常: {type(e).__name__}: {e}")
            logger.error(f"程序异常退出: {traceback.format_exc()}")
            print("5秒后自动重启...")
            time.sleep(5)
            continue
