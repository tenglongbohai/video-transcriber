# -*- coding: utf-8 -*-
"""
导出Word模块 - 生成Word文档
"""
from datetime import datetime

# 尝试导入 docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordExporter:
    """Word 导出"""

    def create_document(self, paragraphs, output_path, title="转录文档"):
        """创建文档"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx: pip install python-docx"

        try:
            doc = Document()
            doc.core_properties.title = title

            heading = doc.add_heading(title, 0)
            heading.alignment = 1

            doc.add_paragraph(f"创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    def create_initial_document(self, output_path, title="转录文档"):
        """创建初始文档（带标题）"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx: pip install python-docx"

        try:
            doc = Document()
            doc.core_properties.title = title

            heading = doc.add_heading(title, 0)
            heading.alignment = 1

            doc.add_paragraph(f"创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"[自动保存模式]")

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    def append_paragraphs(self, paragraphs, output_path):
        """追加段落到已存在的文档"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx: pip install python-docx"

        try:
            doc = Document(output_path)

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)


if __name__ == "__main__":
    # 测试Word导出模块
    exporter = WordExporter()
    if DOCX_AVAILABLE:
        print("Word导出模块测试成功 - python-docx 已安装")
    else:
        print("Word导出模块测试成功 - 请安装 python-docx")
