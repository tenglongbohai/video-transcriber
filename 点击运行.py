#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
视频转文字工具 - Web 服务器
"""

import os
import sys
import json
import time
import uuid
import queue
import threading
import traceback
import subprocess
import logging
from datetime import datetime
from pathlib import Path
from functools import wraps

# 启动时立即初始化日志，确保所有错误都能记录
LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'web_server.log')
try:
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(LOG_FILE, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
except Exception as e:
    # 如果日志初始化失败，写入文件
    with open(LOG_FILE, 'w', encoding='utf-8') as f:
        f.write(f"[ERROR] 日志初始化失败: {e}\n")

logger = logging.getLogger(__name__)

# 启动时自动清理旧日志
def cleanup_old_logs():
    """启动时清理旧日志"""
    try:
        # 清理任务日志
        task_log = os.path.join(os.path.dirname(__file__), "输出文件", "task.log")
        if os.path.exists(task_log):
            with open(task_log, 'w', encoding='utf-8') as f:
                f.write("")
        
        # 清理进度文件
        progress_file = os.path.join(os.path.dirname(__file__), '输出文件', 'task_progress.json')
        if os.path.exists(progress_file):
            os.remove(progress_file)
        
        logger.info("旧日志已清理")
    except Exception as e:
        logger.warning(f"清理日志失败: {e}")

cleanup_old_logs()

# 自动安装依赖
def install_deps():
    deps = {
        "flask": "flask",
        "psutil": "psutil",
        "faster-whisper": "faster-whisper",
        "python-docx": "python-docx",
    }
    for module, package in deps.items():
        try:
            __import__(module)
        except ImportError:
            print(f"正在安装 {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package, "-q"])
    # 检查 ffmpeg
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
    except:
        print("\n" + "!" * 50)
        print("警告: ffmpeg 未安装!")
        print("转录功能可能不可用")
        print("如需转录，请手动安装: winget install Gyan.FFmpeg")
        print("或从 https://ffmpeg.org 下载")
        print("!" * 50)

install_deps()

logger.info("程序启动成功")

from flask import (
    Flask, request, send_file, jsonify,
    render_template, Response
)
import psutil

# 尝试导入可选依赖
try:
    from faster_whisper import WhisperModel
    FASTER_WHISPER_AVAILABLE = True
except ImportError:
    FASTER_WHISPER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# 创建 Flask 应用
app = Flask(__name__, static_folder='web', template_folder='web')
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024 * 1024  # 10GB max
app.config['UPLOAD_FOLDER'] = '输出文件'

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 全局状态
class TaskState:
    def __init__(self):
        self.is_running = False
        self.is_paused = False
        self.is_stopped = False
        self.progress = 0
        self.stage = ""
        self.current_chunk = 0
        self.total_chunks = 0
        self.time_left = ""
        self.last_update = time.time()
        self.current_task_id = None

state = TaskState()

# SSE 消息队列
message_queue = queue.Queue()

# 多任务队列
class TaskQueue:
    def __init__(self):
        self.tasks = []
        self.lock = threading.Lock()

    def add_task(self, task_info):
        with self.lock:
            self.tasks.append(task_info)

    def get_next_task(self):
        with self.lock:
            for task in self.tasks:
                if task.get('status') == 'waiting':
                    return task
            return None

    def update_task(self, task_id, **kwargs):
        with self.lock:
            for task in self.tasks:
                if task.get('id') == task_id:
                    task.update(kwargs)
                    return True
            return False

    def remove_task(self, task_id):
        with self.lock:
            self.tasks = [t for t in self.tasks if t.get('id') != task_id]

    def get_task(self, task_id):
        with self.lock:
            for task in self.tasks:
                if task.get('id') == task_id:
                    return task
            return None

    def has_waiting_tasks(self):
        with self.lock:
            return any(t.get('status') == 'waiting' for t in self.tasks)

task_queue_manager = TaskQueue()




def send_sse(data, task_id=None):
    """发送 Server-Sent Events"""
    try:
        # 添加 task_id 到数据中
        if task_id:
            data['task_id'] = task_id
        message_queue.put_nowait(json.dumps(data))
    except:
        pass
    # 同时写入任务日志文件
    if data.get("type") == "log":
        level = data.get("level", "info")
        msg = data.get("message", "")
        log_file = os.path.join(os.path.dirname(__file__), "输出文件", "task.log")
        try:
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(f"[{level.upper()}] {msg}\n")
        except:
            pass


def sse_generator():
    """SSE 流生成器"""
    import time as time_module
    while True:  # 始终保持连接
        try:
            data = message_queue.get(timeout=5)
            yield f"data: {data}\n\n"
        except queue.Empty:
            # 发送心跳保持连接
            yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"


# ============================================================================
# 核心模块
# ============================================================================

class AudioExtractor:
    """FFmpeg 音频提取器"""

    def __init__(self, ffmpeg_path="ffmpeg"):
        self.ffmpeg_path = ffmpeg_path

    def extract_audio(self, video_path, output_path, sample_rate=16000, progress_callback=None):
        """提取音频"""
        cmd = [
            self.ffmpeg_path, "-y", "-i", video_path,
            "-vn", "-acodec", "pcm_s16le",
            "-ar", str(sample_rate), "-ac", "1",
            "-loglevel", "error", output_path
        ]

        try:
            process = subprocess.Popen(
                cmd, stdout=subprocess.PIPE,
                stderr=subprocess.PIPE, universal_newlines=True
            )

            duration = self._get_duration(video_path)

            while True:
                line = process.stderr.readline()
                if not line and process.poll() is not None:
                    break
                if line and "time=" in line and progress_callback:
                    current_time = self._parse_time(line)
                    if duration > 0:
                        progress = min(95, int(current_time / duration * 100))
                        progress_callback(progress)

            if process.returncode == 0:
                if progress_callback:
                    progress_callback(100)
                return True, "完成"
            else:
                return False, process.stderr.read()

        except Exception as e:
            return False, str(e)

    def _get_duration(self, media_path):
        """获取媒体文件时长（支持视频和音频）"""
        cmd = [
            self.ffmpeg_path, "-i", media_path,
            "-f", "null", "-"
        ]
        try:
            result = subprocess.run(
                cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE,
                timeout=30, text=True
            )
            stderr = result.stderr

            # 尝试解析多种格式
            # 格式1: Duration: 00:05:30.50
            for line in stderr.split('\n'):
                if "Duration:" in line:
                    try:
                        d = line.split("Duration:")[1].split(",")[0].strip()
                        h, m, s = d.split(":")
                        duration = int(h) * 3600 + int(m) * 60 + float(s)
                        if duration > 0:
                            return duration
                    except:
                        pass

            # 格式2: 直接解析 time= 字段
            for line in stderr.split('\n'):
                if "time=" in line:
                    try:
                        t = line.split("time=")[1].split(" ")[0].strip()
                        h, m, s = t.split(":")
                        duration = int(h) * 3600 + int(m) * 60 + float(s)
                        if duration > 0:
                            return duration
                    except:
                        pass

        except:
            pass

        # 备用方法：使用 ffprobe
        try:
            ffprobe_path = self.ffmpeg_path.replace("ffmpeg", "ffprobe")
            cmd2 = [
                ffprobe_path, "-v", "error", "-show_entries",
                "format=duration", "-of", "default=noprint_wrappers=1:nokey=1",
                media_path
            ]
            result2 = subprocess.run(cmd2, stderr=subprocess.PIPE, stdout=subprocess.PIPE, timeout=30, text=True)
            if result2.stdout.strip():
                return float(result2.stdout.strip())
        except:
            pass

        return 0

    def _parse_time(self, line):
        try:
            t = line.split("time=")[1].split(" ")[0]
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        except:
            return 0

    def cut_segment(self, audio_path, output_path, start_sec, duration_sec, progress_callback=None):
        """切割音频片段"""
        cmd = [
            self.ffmpeg_path, "-y",
            "-i", audio_path,
            "-ss", str(start_sec),
            "-t", str(duration_sec),
            "-acodec", "pcm_s16le",
            "-ar", "16000", "-ac", "1",
            "-loglevel", "error",
            output_path
        ]
        try:
            result = subprocess.run(cmd, capture_output=True, timeout=duration_sec + 30)
            if result.returncode == 0:
                if progress_callback:
                    progress_callback(100)
                return True, "完成"
            else:
                return False, result.stderr.decode('utf-8', errors='ignore')
        except subprocess.TimeoutExpired:
            return False, "切割超时"
        except Exception as e:
            return False, str(e)


class Transcriber:
    """Whisper 转录器"""

    def __init__(self, model_size="small", compute_type="int8_float16"):
        self.model_size = model_size
        self.compute_type = compute_type
        self.model = None

    def load_model(self):
        """加载模型"""
        if not FASTER_WHISPER_AVAILABLE:
            return False, "请安装 faster-whisper"

        try:
            # 使用 float32 稳定兼容
            compute = "float32"

            self.model = WhisperModel(
                self.model_size,
                device="cuda",
                compute_type=compute
            )
            return True, "模型加载成功"
        except Exception as e:
            return False, str(e)

    def transcribe(self, audio_path, chunk_start=0, chunk_end=None):
        """转录"""
        if self.model is None:
            success, msg = self.load_model()
            if not success:
                return False, msg, ""

        try:
            kwargs = {
                "audio": audio_path,
                "vad_filter": False,  # 关闭VAD，避免时间戳问题
                "language": "zh",
            }

            segments, info = self.model.transcribe(**kwargs)

            text_parts = []
            for segment in segments:
                seg_start = segment.start
                # 过滤时间范围
                if chunk_start > 0 and seg_start < chunk_start:
                    continue
                if chunk_end and seg_start >= chunk_end:
                    break
                text = segment.text.strip()
                if text:
                    text_parts.append(text)

            return True, "完成", " ".join(text_parts)
        except Exception as e:
            return False, str(e), ""


class PunctuationAdder:
    """标点处理"""

    def add_punctuation(self, text):
        """添加标点 - 智能分析文本结构添加标点"""
        if not text:
            return text

        # 清理文本
        text = text.strip()

        # 检测是否有标点，如果没有就智能添加
        has_punctuation = any(c in '。！？.?!' for c in text)

        if not has_punctuation:
            # 智能添加标点：按长度和常见停顿点分段
            import re
            result = []
            current = ""

            # 按空格或换行分割成句子单位
            words = re.split(r'[\s\n]+', text)

            for word in words:
                if not word.strip():
                    continue
                current += word + " "

                # 根据当前长度添加标点
                if len(current) >= 80:  # 长句用句号
                    # 找合适的断点（常见连词、转折词）
                    for sep in ['，', '的', '了', '是', '在', '和', '也', '就', '但']:
                        if sep in current[:-1]:
                            idx = current.rfind(sep)
                            if idx > 20:  # 确保前面有足够内容
                                result.append(current[:idx+1].strip() + "。")
                                current = current[idx+1:].strip()
                                if current:
                                    current += " "
                                break
                    else:
                        # 没找到断点，直接截断
                        if current.strip():
                            result.append(current.strip() + "。")
                            current = ""
                elif len(current) >= 50 and current.rstrip().endswith(('，', '、')):
                    # 中等长度，遇到逗号可以考虑结束
                    pass

            # 处理剩余内容
            if current.strip():
                if len(current.strip()) > 15:
                    result.append(current.strip() + "。")
                else:
                    result.append(current.strip())

            # 合并结果
            if result:
                return " ".join(result)

        # 已有标点，简单整理
        sentences = []
        current = ""
        for char in text:
            current += char
            if char in '。！？.?!':
                sentences.append(current.strip())
                current = ""

        if current.strip():
            sentences.append(current.strip())

        return " ".join(sentences) if sentences else text

    def split_paragraphs(self, text, max_length=500):
        """分段"""
        import re
        paragraphs = []
        current = ""

        parts = re.split(r'[。！？\?!]+', text)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            if len(current) + len(part) <= max_length:
                current += part + "。 "
            else:
                if current.strip():
                    paragraphs.append(current.strip())
                current = part + "。 "

        if current.strip():
            paragraphs.append(current.strip())

        return paragraphs if paragraphs else [text]


class MiniMaxAPI:
    """MiniMax API - Anthropic 兼容接口"""

    BASE_URL = "https://api.minimaxi.com/anthropic"
    MODELS = ["MiniMax-M2.5", "MiniMax-M2.7", "MiniMax-M2.1", "MiniMax-M2"]

    def __init__(self, api_key=None, model="MiniMax-M2.5"):
        self.api_key = api_key  # 必须从外部传入
        self.model = model
        self.headers = {
            "x-api-key": self.api_key,
            "Content-Type": "application/json",
            "anthropic-version": "2023-06-01"
        }

    def polish_text(self, text, log_callback=None):
        """润色文本"""
        def log(msg, level="info"):
            print(msg)
            if log_callback:
                log_callback(msg, level)

        if not text or not text.strip():
            return text

        if not self.api_key:
            log("未配置 MiniMax API Key，跳过润色", "warning")
            return text

        if not ANTHROPIC_AVAILABLE:
            log("请安装 anthropic SDK: pip install anthropic", "error")
            return text

        prompt = f"""你是一个专业的文本润色助手。请对以下转录文本进行处理：
1. 首先分析内容，生成一个简洁的视频大纲（用数字列表列出主要话题）
2. 然后对原文进行润色：添加标点、修正错误、保持口语风格
3. **必须使用简体中文输出**，不要使用繁体字
4. 严格按照以下格式输出，不要添加任何解释：

---
**视频大纲：**
:[根据内容生成的大纲，每行一个话题]

**正文：**
:[润色后的正文]
---

原文：
{text}"""

        max_retries = 5
        retry_delay = 5

        for attempt in range(max_retries):
            try:
                client = anthropic.Anthropic(
                    base_url=self.BASE_URL,
                    api_key=self.api_key
                )

                if attempt > 0:
                    log(f"正在调用 MiniMax API（第 {attempt + 1} 次重试）...")
                else:
                    log("正在调用 MiniMax API...")

                message = client.messages.create(
                    model=self.model,
                    max_tokens=4096,
                    temperature=1.0,
                    system="你是一个专业的文本润色助手。直接输出润色结果，不要添加解释或说明。",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": prompt
                                }
                            ]
                        }
                    ]
                )

                for block in message.content:
                    if block.type == "text":
                        raw_result = block.text.strip()
                        log(f"润色完成，返回 {len(raw_result)} 字符", "info")
                        
                        # 提取正文部分（去掉视频大纲）
                        # 格式: --- **视频大纲：** ... **正文：** ...
                        if "**正文：**" in raw_result or "**正文:**" in raw_result:
                            lines = raw_result.split('\n')
                            in_body = False
                            body_lines = []
                            for line in lines:
                                if '**正文' in line:
                                    in_body = True
                                    # 不跳过，继续处理后续行
                                    continue
                                if in_body:
                                    # 跳过空行
                                    if not line.strip():
                                        continue
                                    body_lines.append(line.strip())
                            if body_lines:
                                return '\n'.join(body_lines)
                        
                        # 如果没有标准格式，直接返回原文本（让后续标点处理）
                        return raw_result

                log("润色返回空结果", "warning")
                return text

            except Exception as e:
                error_str = str(e)
                if "overloaded_error" in error_str or "529" in error_str:
                    if attempt < max_retries - 1:
                        log(f"MiniMax 服务器繁忙，{retry_delay}秒后重试...", "warning")
                        import time as time_module
                        time_module.sleep(retry_delay)
                        retry_delay *= 2
                    else:
                        log(f"MiniMax API 重试 {max_retries} 次仍失败: {e}", "error")
                        return text
                else:
                    log(f"MiniMax API 错误: {e}", "error")
                    return text


class WordExporter:
    """Word 导出"""

    def create_document(self, paragraphs, output_path, title="转录文档"):
        """创建文档"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx"

        try:
            doc = Document()
            doc.core_properties.title = title

            heading = doc.add_heading(title, 0)
            heading.alignment = 1

            doc.add_paragraph(f"创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    def create_initial_document(self, output_path, title="转录文档"):
        """创建初始文档（带标题）"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx"

        try:
            doc = Document()
            doc.core_properties.title = title

            heading = doc.add_heading(title, 0)
            heading.alignment = 1

            doc.add_paragraph(f"创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"[自动保存模式]")

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    def append_paragraphs(self, paragraphs, output_path):
        """追加段落到已存在的文档"""
        if not DOCX_AVAILABLE:
            return False, "请安装 python-docx"

        try:
            doc = Document(output_path)

            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5

            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)


# ============================================================================
# 转录任务
# ============================================================================

def load_progress(video_path, output_dir):
    """加载断点续录进度"""
    progress_file = os.path.join(output_dir, f".progress_{uuid.uuid4().hex[:8]}.json")
    if os.path.exists(progress_file):
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return None

def save_progress(progress_file, data):
    """保存进度"""
    try:
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)
    except:
        pass

def run_transcription(video_path, output_dir, config, task_id=None):
    """执行转录任务"""
    global state
    import time
    task_start_time = time.time()

    # 清空任务日志
    log_file = os.path.join(os.path.dirname(__file__), "输出文件", "task.log")
    try:
        with open(log_file, "w", encoding="utf-8") as f:
            f.write("")
    except:
        pass

    state.is_running = True
    state.is_paused = False
    state.is_stopped = False
    state.progress = 0
    state.stage = "初始化"
    state.current_chunk = 0
    state.total_chunks = 0
    state.current_task_id = task_id

    # 定义带 task_id 的 send_sse 包装，同时保存进度用于轮询恢复
    def send_task_sse(data):
        send_sse(data, task_id)
        # 保存关键进度到内存/文件
        if data.get("type") in ["progress", "log", "complete", "error"]:
            save_task_progress(task_id, data)

    temp_files = []
    progress_file = None

    try:
        # 记录任务开始信息
        send_task_sse({"type": "log", "level": "info", "message": f"========== 任务开始 =========="})
        send_task_sse({"type": "log", "level": "info", "message": f"视频路径: {video_path}"})
        send_task_sse({"type": "log", "level": "info", "message": f"视频大小: {os.path.getsize(video_path) / 1024 / 1024:.2f} MB"})
        send_task_sse({"type": "log", "level": "info", "message": f"Whisper模型: {config['model']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"分段时长: {config['chunk_duration']} 分钟"})
        send_task_sse({"type": "log", "level": "info", "message": f"启用标点: {config['punctuation']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"启用润色: {config['polish']}"})
        send_task_sse({"type": "log", "level": "info", "message": f"输出目录: {output_dir}"})
        send_task_sse({"type": "log", "level": "info", "message": "=" * 40})

        # 初始化
        send_task_sse({"type": "log", "level": "info", "message": "开始初始化组件..."})
        audio_extractor = AudioExtractor()
        transcriber = Transcriber(config["model"])
        punctuation_adder = PunctuationAdder()
        word_exporter = WordExporter()

        # 输出文件
        output_name = config.get("output_name", "转录文档")
        output_path = os.path.join(output_dir, output_name + ".docx")

        # 进度文件
        progress_file = os.path.join(output_dir, f".{output_name}_progress.json")

        # 阶段1: 音频提取
        send_task_sse({"type": "log", "level": "info", "message": "正在提取音频..."})
        state.stage = "音频提取"

        temp_audio = os.path.join(output_dir, f".temp_{uuid.uuid4().hex}.wav")
        temp_files.append(temp_audio)

        success, msg = audio_extractor.extract_audio(
            video_path, temp_audio,
            progress_callback=lambda p: send_task_sse({
                "type": "progress",
                "progress": int(p * 0.05),
                "stage": f"音频提取 {p}%"
            })
        )

        if not success:
            raise Exception(f"音频提取失败: {msg}")

        send_task_sse({"type": "log", "level": "info", "message": "音频提取完成"})

        # 加载模型
        send_task_sse({"type": "log", "level": "info", "message": "加载 Whisper 模型..."})
        success, msg = transcriber.load_model()
        if not success:
            raise Exception(msg)

        send_task_sse({"type": "log", "level": "info", "message": f"模型加载完成: {transcriber.model_size}"})

        # 获取总时长
        total_duration = audio_extractor._get_duration(temp_audio)
        chunk_duration = int(config.get("chunk_duration", 5)) * 60
        total_chunks = max(1, int(total_duration / chunk_duration) + (1 if total_duration % chunk_duration > 0 else 0))
        state.total_chunks = total_chunks

        send_task_sse({"type": "log", "level": "info", "message": f"音频时长: {total_duration:.0f}秒, 分为 {total_chunks} 个片段"})

        # 检查断点续录（通过 output_name 和原始视频文件匹配）
        start_chunk_idx = 0
        completed_texts = []
        if os.path.exists(progress_file):
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    saved = json.load(f)
                    saved_video = saved.get('video_path', '')
                    
                    # 转换相对路径为绝对路径
                    if not os.path.isabs(saved_video):
                        saved_video_abs = os.path.join(os.path.dirname(__file__), saved_video)
                    else:
                        saved_video_abs = saved_video
                    
                    # 必须原始视频文件还存在，且文件名匹配，才续传
                    video_exists = os.path.exists(saved_video_abs)
                    
                    # 比较文件名（去掉uuid部分）
                    original_name = os.path.basename(video_path)
                    saved_original_name = os.path.basename(saved_video)
                    if '_' in original_name and '_' in saved_original_name:
                        orig_short = '_'.join(original_name.split('_')[1:])
                        saved_short = '_'.join(saved_original_name.split('_')[1:])
                        name_match = (orig_short == saved_short)
                    else:
                        name_match = (original_name == saved_original_name)
                    
                    # 只有视频文件存在且文件名匹配时才续传
                    if video_exists and name_match and saved.get('total_chunks') == total_chunks:
                        start_chunk_idx = saved.get('last_completed_chunk', 0) + 1
                        completed_texts = saved.get('texts', [])
                        send_task_sse({"type": "log", "level": "info", "message": f"检测到进度文件，从片段 {start_chunk_idx + 1} 继续（已完成 {len(completed_texts)} 个片段）"})
                    elif not video_exists:
                        # 原始视频不存在，视为新任务
                        send_task_sse({"type": "log", "level": "info", "message": "未检测到之前的视频文件，开始新任务"})
            except:
                pass

        # 阶段2: FFmpeg预切割 + 分段转录
        start_time = time.time()
        base_progress = 5

        def calc_time_left(current, total):
            """计算预估剩余时间"""
            elapsed = time.time() - start_time
            if current == 0:
                return ""
            avg_time_per_chunk = elapsed / current
            remaining = int(avg_time_per_chunk * (total - current))
            if remaining >= 60:
                return f"{remaining // 60}分{remaining % 60}秒"
            return f"{remaining}秒"

        for chunk_idx in range(start_chunk_idx, total_chunks):
            # 检查停止
            if state.is_stopped:
                send_task_sse({"type": "log", "level": "warning", "message": "任务已停止"})
                # 保存当前进度
                save_progress(progress_file, {
                    'video_path': video_path,
                    'total_chunks': total_chunks,
                    'last_completed_chunk': chunk_idx - 1,
                    'texts': completed_texts
                })
                send_task_sse({"type": "log", "level": "info", "message": f"进度已保存，可点击继续从片段 {chunk_idx} 恢复"})
                break

            # 等待暂停
            while state.is_paused and not state.is_stopped:
                time.sleep(0.5)

            chunk_start = chunk_idx * chunk_duration
            chunk_end = min(chunk_start + chunk_duration, total_duration)
            actual_duration = chunk_end - chunk_start
            state.current_chunk = chunk_idx + 1
            state.stage = f"转录片段 {chunk_idx + 1}/{total_chunks}"

            # FFmpeg 切割这段音频
            segment_file = os.path.join(output_dir, f".segment_{chunk_idx:03d}.wav")
            temp_files.append(segment_file)

            send_task_sse({"type": "log", "level": "info", "message": f"切割片段 {chunk_idx + 1}/{total_chunks}..."})
            success, msg = audio_extractor.cut_segment(
                temp_audio, segment_file, chunk_start, actual_duration
            )
            if not success:
                send_task_sse({"type": "log", "level": "warning", "message": f"片段 {chunk_idx + 1} 切割失败: {msg}，跳过"})
                continue

            # 进度更新
            chunk_progress = (chunk_idx - start_chunk_idx + 1) / (total_chunks - start_chunk_idx) * 60
            time_left = calc_time_left(chunk_idx - start_chunk_idx + 1, total_chunks - start_chunk_idx)
            send_task_sse({
                "type": "progress",
                "progress": base_progress + int(chunk_progress),
                "stage": state.stage,
                "current": chunk_idx + 1,
                "total": total_chunks,
                "time_left": time_left
            })

            send_task_sse({"type": "log", "level": "info", "message": f"开始转录片段 {chunk_idx + 1}/{total_chunks} ({chunk_start:.0f}-{chunk_end:.0f}秒)"})

            # 转录这段音频（现在每段时间长固定，不会越往后越慢）
            success, msg, text = transcriber.transcribe(
                segment_file,
                chunk_start=0,
                chunk_end=None
            )

            if success and text:
                # 每段转录完立即保存到Word
                paragraphs = punctuation_adder.split_paragraphs(text, 500)

                # 如果是第一个片段，创建文档
                if chunk_idx == start_chunk_idx and not completed_texts:
                    success, msg = word_exporter.create_initial_document(
                        output_path,
                        title=f"视频转录 - {output_name}"
                    )
                    if not success:
                        send_task_sse({"type": "log", "level": "warning", "message": f"文档创建失败: {msg}"})

                # 追加到文档
                success, msg = word_exporter.append_paragraphs(paragraphs, output_path)
                if success:
                    send_task_sse({"type": "log", "level": "info", "message": f"片段 {chunk_idx + 1} 已保存 ({len(text)} 字符)"})
                else:
                    send_task_sse({"type": "log", "level": "warning", "message": f"文档追加失败: {msg}"})

                completed_texts.append(text)

                # 保存进度
                save_progress(progress_file, {
                    'video_path': video_path,
                    'total_chunks': total_chunks,
                    'last_completed_chunk': chunk_idx,
                    'texts': completed_texts
                })
            else:
                send_task_sse({"type": "log", "level": "warning", "message": f"片段 {chunk_idx + 1} 转录失败: {msg}"})

            # 删除临时片段文件释放空间
            try:
                if os.path.exists(segment_file):
                    os.remove(segment_file)
                    temp_files.remove(segment_file)
            except:
                pass

        # 保留进度文件（不清除），支持后续续传或重新处理
        # if os.path.exists(progress_file) and not state.is_stopped:
        #     try:
        #         os.remove(progress_file)
        #     except:
        #         pass

        send_task_sse({"type": "log", "level": "info", "message": f"转录完成，共 {len(completed_texts)} 个片段"})

        if not completed_texts:
            raise Exception("没有成功转录任何内容")

        # 阶段3: 润色或标点
        full_text = " ".join(completed_texts)
        send_task_sse({"type": "log", "level": "info", "message": f"原始文本: {len(full_text)} 字符"})

        state.stage = "后处理"

        if config.get("polish", False):
            api_key = config.get("api_key", "").strip()
            if not api_key:
                send_task_sse({"type": "log", "level": "warning", "message": "未提供 MiniMax API Key，无法润色"})
                full_text = punctuation_adder.add_punctuation(full_text)
                send_task_sse({"type": "progress", "progress": 75, "stage": "添加标点", "time_left": "即将完成"})
            else:
                send_task_sse({"type": "log", "level": "info", "message": "正在使用 MiniMax 润色..."})
                try:
                    minimax = MiniMaxAPI(api_key, model=config.get("minimax_model", "MiniMax-M2.5"))
                    segments = punctuation_adder.split_paragraphs(full_text, 800)
                    polish_start = time.time()
                    polished = []

                    def polish_log(msg, level="info"):
                        send_task_sse({"type": "log", "level": level, "message": msg})

                    for i, seg in enumerate(segments):
                        polished.append(minimax.polish_text(seg, log_callback=polish_log))
                        elapsed = time.time() - polish_start
                        remaining = int(elapsed / (i + 1) * (len(segments) - i - 1))
                        time_str = f"{remaining}秒" if remaining < 60 else f"{remaining // 60}分{remaining % 60}秒"
                        send_task_sse({
                            "type": "progress",
                            "progress": 60 + int((i + 1) / len(segments) * 20),
                            "stage": f"润色中 {i + 1}/{len(segments)}",
                            "time_left": time_str
                        })
                    full_text = " ".join(polished)
                    send_task_sse({"type": "log", "level": "info", "message": f"润色完成: {len(full_text)} 字符"})
                    
                    # 检查润色结果是否有标点，如果没有则添加标点
                    has_punct = any(c in '。！？.?!' for c in full_text)
                    if not has_punct:
                        send_task_sse({"type": "log", "level": "info", "message": "润色内容无标点，正在添加标点..."})
                        full_text = punctuation_adder.add_punctuation(full_text)
                        send_task_sse({"type": "log", "level": "info", "message": f"标点添加完成: {len(full_text)} 字符"})
                    
                    send_task_sse({"type": "progress", "progress": 80, "stage": "后处理", "time_left": "即将完成"})
                except Exception as e:
                    send_task_sse({"type": "log", "level": "warning", "message": f"润色失败: {e}，使用原始文本"})
                    full_text = punctuation_adder.add_punctuation(full_text)
                    send_task_sse({"type": "progress", "progress": 75, "stage": "添加标点", "time_left": "即将完成"})
        else:
            full_text = punctuation_adder.add_punctuation(full_text)
            send_task_sse({"type": "progress", "progress": 75, "stage": "添加标点", "time_left": "即将完成"})

        send_task_sse({"type": "log", "level": "info", "message": f"标点处理后: {len(full_text)} 字符"})

        # 完成
        send_task_sse({"type": "progress", "progress": 100, "stage": "完成", "time_left": "0秒"})
        send_task_sse({
            "type": "complete",
            "output_path": output_path,
            "output_name": output_name,
            "memory": f"{psutil.virtual_memory().percent:.0f}%"
        })

        state.progress = 100
        state.stage = "完成"
        total_time = time.time() - task_start_time
        send_task_sse({"type": "log", "level": "info", "message": "=" * 40})
        send_task_sse({"type": "log", "level": "info", "message": f"任务完成！总耗时: {total_time:.1f}秒 ({total_time/60:.1f}分钟)"})
        send_task_sse({"type": "log", "level": "info", "message": f"输出文件: {output_path}"})
        if os.path.exists(output_path):
            send_task_sse({"type": "log", "level": "info", "message": f"文件大小: {os.path.getsize(output_path)/1024:.1f} KB"})
        send_task_sse({"type": "log", "level": "info", "message": f"========== 任务结束 =========="})

    except Exception as e:
        send_task_sse({"type": "error", "message": str(e)})
        send_task_sse({"type": "log", "level": "error", "message": f"错误: {traceback.format_exc()}"})
        send_task_sse({"type": "log", "level": "error", "message": f"任务失败！耗时: {time.time() - task_start_time:.1f}秒"})
        send_task_sse({"type": "log", "level": "error", "message": "=" * 40})

    finally:
        # 清理临时文件
        for f in temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except:
                pass

        state.is_running = False


# ============================================================================
# Flask 路由
# ============================================================================

@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/transcribe', methods=['POST'])
def transcribe():
    """开始转录"""
    if state.is_running:
        return jsonify({"error": "已有任务在运行"}), 400

    video = request.files.get('video')
    if not video:
        return jsonify({"error": "未上传视频文件"}), 400

    # 保存上传的视频
    video_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex}_{video.filename}")
    video.save(video_path)

    # 获取视频所在目录作为输出目录
    video_dir = os.path.dirname(video_path)

    # 获取 task_id（用于多任务队列）
    task_id = request.form.get('task_id', str(uuid.uuid4().hex[:8]))

    # 获取配置
    config = {
        "output_name": request.form.get('output_name', '转录文档'),
        "model": request.form.get('model', 'small'),
        "chunk_duration": int(request.form.get('chunk_duration', 5)),
        "punctuation": request.form.get('punctuation', 'true').lower() == 'true',
        "polish": request.form.get('polish', 'false').lower() == 'true',
        "api_key": request.form.get('api_key', ''),
        "minimax_model": request.form.get('minimax_model', 'MiniMax-M2.5'),
        "video_path": video_path,
        "output_dir": video_dir  # 保存到视频所在目录
    }

    # 启动后台任务
    thread = threading.Thread(target=run_transcription, args=(video_path, app.config['UPLOAD_FOLDER'], config, task_id))
    thread.daemon = False  # 非守护线程，确保任务完成
    thread.start()

    return jsonify({"status": "started", "video_path": video_path, "task_id": task_id})


@app.route('/api/resume', methods=['POST'])
def resume():
    """继续转录（断点续录）"""
    if state.is_running:
        return jsonify({"error": "已有任务在运行"}), 400

    video_path = request.form.get('video_path')
    output_name = request.form.get('output_name', '转录文档')
    if not video_path or not os.path.exists(video_path):
        return jsonify({"error": "视频文件不存在"}), 400

    output_dir = os.path.dirname(video_path)
    progress_file = os.path.join(output_dir, f".{output_name}_progress.json")
    if not os.path.exists(progress_file):
        return jsonify({"error": "没有可恢复的进度"}), 400

    # 加载进度
    try:
        with open(progress_file, 'r', encoding='utf-8') as f:
            saved = json.load(f)
    except:
        return jsonify({"error": "进度文件读取失败"}), 400

    # 检查视频是否匹配
    if saved.get('video_path') != video_path:
        return jsonify({"error": "视频文件不匹配，无法恢复"}), 400

    # 获取 task_id（用于多任务队列）
    task_id = request.form.get('task_id', str(uuid.uuid4().hex[:8]))

    config = {
        "output_name": output_name,
        "model": request.form.get('model', 'small'),
        "chunk_duration": int(request.form.get('chunk_duration', 5)),
        "punctuation": request.form.get('punctuation', 'true').lower() == 'true',
        "polish": request.form.get('polish', 'false').lower() == 'true',
        "api_key": request.form.get('api_key', ''),
        "minimax_model": request.form.get('minimax_model', 'MiniMax-M2.5'),
        "video_path": video_path,
        "output_dir": output_dir,
        "resume_from": saved.get('last_completed_chunk', -1) + 1
    }

    # 启动后台任务
    thread = threading.Thread(target=run_transcription, args=(video_path, output_dir, config, task_id))
    thread.daemon = False  # 非守护线程，确保任务完成
    thread.start()

    return jsonify({"status": "resumed", "start_chunk": config["resume_from"], "task_id": task_id})


@app.route('/api/progress')
def progress():
    """SSE 进度流"""
    return Response(sse_generator(), mimetype='text/event-stream')


@app.route('/api/pause', methods=['POST'])
def pause():
    """暂停/继续"""
    state.is_paused = not state.is_paused
    return jsonify({"paused": state.is_paused})


@app.route('/api/stop', methods=['POST'])
def stop():
    """停止任务"""
    state.is_stopped = True
    return jsonify({"stopped": True})


@app.route('/api/status')
def status():
    """获取状态"""
    return jsonify({
        "running": state.is_running,
        "paused": state.is_paused,
        "memory": f"{psutil.virtual_memory().percent:.0f}"
    })

# 任务进度存储（用于轮询恢复）
task_progress_store = {}

def save_task_progress(task_id, data):
    """保存任务进度到内存/文件"""
    task_progress_store[task_id] = data
    # 同时写入文件，防止程序重启后丢失
    progress_file = os.path.join(os.path.dirname(__file__), '输出文件', 'task_progress.json')
    try:
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(task_progress_store, f, ensure_ascii=False)
    except:
        pass

def load_task_progress(task_id):
    """加载任务进度"""
    global task_progress_store
    if task_id in task_progress_store:
        return task_progress_store[task_id]
    # 尝试从文件加载
    progress_file = os.path.join(os.path.dirname(__file__), '输出文件', 'task_progress.json')
    if os.path.exists(progress_file):
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                task_progress_store = json.load(f)
                return task_progress_store.get(task_id)
        except:
            pass
    return None

@app.route('/api/poll_progress/<task_id>')
def poll_progress(task_id):
    """轮询任务进度（用于SSE断连后恢复）"""
    progress = load_task_progress(task_id)
    if progress:
        return jsonify(progress)
    return jsonify({"type": "no_data"})


@app.route('/api/clear_logs', methods=['POST'])
def clear_logs():
    """清理日志文件"""
    try:
        # 清理任务日志
        task_log = os.path.join(os.path.dirname(__file__), "输出文件", "task.log")
        if os.path.exists(task_log):
            with open(task_log, 'w', encoding='utf-8') as f:
                f.write("")
        
        # 清理进度文件
        progress_file = os.path.join(os.path.dirname(__file__), '输出文件', 'task_progress.json')
        if os.path.exists(progress_file):
            os.remove(progress_file)
        
        # 清理全局任务进度存储
        global task_progress_store
        task_progress_store = {}
        
        return jsonify({"success": True, "message": "日志已清理"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/download/<path:filename>')
def download(filename):
    """下载文件"""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(file_path):
        file_path = file_path + '.docx'
    return send_file(file_path, as_attachment=True)


@app.route('/api/minimax/test', methods=['POST'])
def minimax_test():
    """测试 MiniMax API 连接并获取模型列表"""
    try:
        data = request.get_json()
        api_key = data.get('api_key', '').strip()

        if not api_key:
            return jsonify({"success": False, "error": "API Key 不能为空"})

        import urllib.request

        # 测试 API 连接 - 直接发送一条简单消息
        url = "https://api.minimaxi.com/anthropic/v1/messages"
        headers = {
            "x-api-key": api_key,
            "Content-Type": "application/json",
            "anthropic-version": "2023-06-01"
        }
        payload = json.dumps({
            "model": "MiniMax-M2.5",
            "max_tokens": 10,
            "messages": [{"role": "user", "content": "hi"}]
        }).encode('utf-8')

        req = urllib.request.Request(url, data=payload, headers=headers)
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                result = json.loads(response.read().decode('utf-8'))
                # 连接成功，返回可用模型列表
                return jsonify({
                    "success": True,
                    "models": [
                        {"id": "MiniMax-M2.7"},
                        {"id": "MiniMax-M2.5"},
                        {"id": "MiniMax-M2.1"}
                    ]
                })
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8') if e.fp else ''
            try:
                error_json = json.loads(error_body)
                error_msg = error_json.get('error', {}).get('message', error_body) if isinstance(error_json.get('error'), dict) else error_json.get('error', error_body)
            except:
                error_msg = error_body if error_body else f"HTTP {e.code}"
            return jsonify({"success": False, "error": error_msg})
        except Exception as e:
            return jsonify({"success": False, "error": str(e)})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# ============================================================================
# 启动
# ============================================================================

if __name__ == '__main__':
    import webbrowser
    import threading
    import logging

    # 关闭 Flask/Werkzeug 日志
    log = logging.getLogger('werkzeug')
    log.setLevel(logging.ERROR)
    app.logger.setLevel(logging.ERROR)

    print("=" * 50)
    print("视频转文字工具 - Web 版")
    print("=" * 50)
    print("正在启动服务器...")
    print("请访问 http://127.0.0.1:5000")
    print("按 Ctrl+C 可关闭服务器")
    print("=" * 50)

    def open_browser():
        import time
        time.sleep(1)
        webbrowser.open('http://localhost:5000')

    threading.Thread(target=open_browser, daemon=True).start()

    # 启动 Flask 服务，自动重启防止崩溃
    while True:
        try:
            app.run(host='0.0.0.0', port=5000, debug=False, threaded=True, use_reloader=False)
        except KeyboardInterrupt:
            print("\n程序已停止")
            break
        except BaseException as e:
            if not isinstance(e, KeyboardInterrupt):
                print(f"程序异常: {e}")
                print("5秒后重新启动...")
                time.sleep(5)
            else:
                print("\n程序已停止")
                break
